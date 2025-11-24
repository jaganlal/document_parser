import re
import sys
import os
import argparse
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

def create_word_table(markdown_input, output_filename='output.docx'):
    # Parse the markdown
    lines = [line.strip() for line in markdown_input.strip().split('\n')]
    
    # Filter out empty lines
    lines = [line for line in lines if line.strip()]
    
    if not lines:
        print("No content found in markdown input.")
        return

    # Identify separator row (usually |---|---|)
    separator_index = -1
    for i, line in enumerate(lines):
        if re.match(r'^\|?[\s\-:|]+\|?$', line) and set(line) & set('-:'):
            separator_index = i
            break
    
    header_rows = []
    data_rows = []
    
    if separator_index != -1:
        header_rows = [line for line in lines[:separator_index]]
        # Only include lines that look like table rows (start with |)
        data_rows = [line for line in lines[separator_index+1:] if line.strip().startswith('|')]
    else:
        # Assume all data if no separator? Or first row header?
        # Standard markdown table usually has separator. 
        # If not found, treat first row as header.
        header_rows = [lines[0]]
        data_rows = lines[1:]

    # Helper to parse a row line into cells
    def parse_row(line):
        # Remove leading/trailing pipes if they exist
        content = line.strip()
        if content.startswith('|'): content = content[1:]
        if content.endswith('|'): content = content[:-1]
        return [cell.strip() for cell in content.split('|')]

    parsed_header_rows = [parse_row(row) for row in header_rows]
    parsed_data_rows = [parse_row(row) for row in data_rows]
    
    # Normalize column count
    # Find max cols
    all_rows = parsed_header_rows + parsed_data_rows
    if not all_rows:
        return
    
    num_cols = max(len(row) for row in all_rows)
    num_rows = len(all_rows)

    # Create Document
    doc = Document()
    doc.add_heading('Converted Table', level=2)

    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Table Grid'

    # Helper to set cell text
    def set_cell_text(cell, text, bold=False):
        cell.text = ""
        paragraph = cell.paragraphs[0]
        parts = text.split('<br>')
        for i, part in enumerate(parts):
            # Split by superscript/subscript markers
            # Matches ^word or _word
            tokens = re.split(r'(\^[a-zA-Z0-9]+|_[a-zA-Z0-9]+)', part)
            
            for token in tokens:
                if not token:
                    continue
                
                run = paragraph.add_run()
                
                if token.startswith('^'):
                    run.text = token[1:]
                    run.font.superscript = True
                elif token.startswith('_'):
                    run.text = token[1:]
                    run.font.subscript = True
                else:
                    run.text = token
                
                if bold:
                    run.bold = True
            
            if i < len(parts) - 1:
                paragraph.add_run().add_break()
                
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Populate Table
    # Headers
    for r, row_data in enumerate(parsed_header_rows):
        for c, text in enumerate(row_data):
            if c < num_cols:
                set_cell_text(table.cell(r, c), text, bold=True)
    
    # Data
    header_offset = len(parsed_header_rows)
    for r, row_data in enumerate(parsed_data_rows):
        row_idx = r + header_offset
        for c, text in enumerate(row_data):
            if c < num_cols:
                set_cell_text(table.cell(row_idx, c), text, bold=False)

    # Generic Merge Logic - DISABLED based on user feedback
    # The aggressive merge logic was causing issues with empty data cells being merged with previous rows.
    # For now, we only merge explicitly if needed, but since we don't have merge info here, we skip it.
    # This fixes the issue where "Day 15" empty cells were merged with "Once Pretest" cells.
    
    # for r in range(num_rows):
    #     for c in range(num_cols):
    #         cell = table.cell(r, c)
    #         text = cell.text.strip()
    #         
    #         if not text:
    #             # Candidate for merge
    #             # Check Up
    #             if r > 0:
    #                 top_cell = table.cell(r-1, c)
    #                 if top_cell.text.strip():
    #                     top_cell.merge(cell)
    #                     continue # Done with this cell
    #             
    #             # Check Left
    #             if c > 0:
    #                 left_cell = table.cell(r, c-1)
    #                 if left_cell.text.strip():
    #                     left_cell.merge(cell)
    #                     continue

    doc.save(output_filename)
    print(f"Successfully created {output_filename}")

def main():
    parser = argparse.ArgumentParser(description="Convert Markdown table to Word Document")
    parser.add_argument('input', help="Markdown string or file path")
    parser.add_argument('--output', default='experimental_design.docx', help="Output file name")
    
    args = parser.parse_args()
    
    markdown_input = args.input
    if os.path.isfile(args.input):
        try:
            with open(args.input, 'r', encoding='utf-8') as f:
                markdown_input = f.read()
        except Exception as e:
            print(f"Error reading file: {e}")
            sys.exit(1)
            
    create_word_table(markdown_input, args.output)

if __name__ == "__main__":
    main()
